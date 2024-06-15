from atexit import register
import json
from operator import index
import string
from xml.dom.minidom import Document
from django.shortcuts import render, HttpResponse
from django.template import *
from numpy import var
from .templates import *
from itertools import zip_longest
import sys
import re

# from flask import Flask, request

sys.path.append("/Users/mac/aibus/lib/python3.10/site-packages")
from docx import Document
import cv2
import easyocr
import PyPDF2
import ssl

ssl._create_default_https_context = ssl._create_unverified_context



def Home(request):
    context = {}
    return render(request, "index.html", context)


def process_image(request):
    
    if request.method == "POST":
        try:
            file_content = request.FILES["image_1"].file.name
            reader = easyocr.Reader(["vi"])
            result = reader.readtext(file_content)
            mat = cv2.imread(file_content)
            boxes = [line[0] for line in result]
            texts = [line[1] for line in result]
            scores = [line[2] for line in result]

            for box in boxes:
                top_left = (int(box[0][0]), int(box[0][1]))
                bottom_right = (int(box[2][0]), int(box[2][1]))
                cv2.rectangle(mat, top_left, bottom_right, (0, 255, 0), 2)
            
            print(texts)

            # Tìm chỉ số của "Họ và tên"
            try:
                number = [item for item in texts if re.match(r'^\d{12}$', item)]
                index_number = number[0]
                
                uppercase_strings = [item for item in texts if item.isupper()]
                
                index_name = uppercase_strings[-1]

                date_pattern = r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{4}\b'

                
                dates_found = re.findall(date_pattern, ' '.join(texts))
                index_date = dates_found[0]

                index_sex = None

                # Duyệt qua danh sách data để tìm "Nữ" hoặc "Nam"
                for item in texts:
                    if item.strip() == "Nữ":
                        index_sex = "Nữ"
                        break
                    elif item.strip() == "Nam":
                        index_sex = "Nam"
                        break
                

                nation = 'Việt Nam'
                
            except ValueError as e:
                # Xử lý trường hợp không tìm thấy mục nào đó
                return HttpResponse(f"Error: {str(e)}", status=500)
            
            
            # return HttpResponse(full_name)
            result_array = [index_number +'_', index_name +'_', index_date +'_', index_sex +'_', nation]
            print(result_array)
            # print(result_array)
            return HttpResponse(result_array)
        

        except Exception as e:
            print(f"Exception: {e}")
            return HttpResponse("Internal Server Error", status=500)
        
    return HttpResponse("Method not allowed", status=405)

def Login(request):
    context = {}
    return render(request, "login.html", context)


def Createprofile(request):
    context = {}
    return render(request, "createprofile.html", context)


def Success(request):
    context = {}
    return render(request, "success.html", context)


def Error(request):
    context = {}
    return render(request, "error.html", context)


def Recruitment(request):
    context = {}
    return render(request, "recruitment.html", context)


# @register
# def convert_doc(request):
#     document = docx.Document("home/test.docx")
#     for para in document.paragraphs:
#         print(para.text)
#     # document.save("test.docx")
#     return document
def handle_fill_data_request(request):
    try:
        # Lấy dữ liệu từ yêu cầu POST
        json_data = json.loads(request.body)
        # Dữ liệu để điền vào DOCX từ json_data
        data = {
            '{{province}}': json_data.get('province', ''),
            '{{date}}': json_data.get('date', ''),
            '{{month}}': json_data.get('month', ''),
            '{{year}}': json_data.get('year', ''),
            '{{name}}': json_data.get('name', ''),
            '{{table_1}}': json_data.get('table_1', ''),
            '{{company_name}}': json_data.get('company_name', ''),
            '{{company_name_en}}': json_data.get('company_name_en', ''),
            '{{company_name_short}}': json_data.get('company_name_short', ''),
            '{{company_address}}': json_data.get('company_address', ''),
            '{{company_ward}}': json_data.get('company_ward', ''),
            '{{company_district}}': json_data.get('company_district', ''),
            '{{company_province}}': json_data.get('company_province', ''),
            '{{company_phone}}': json_data.get('company_phone', ''),
            '{{conpany_fax}}': json_data.get('conpany_fax', ''),
            '{{company_email}}': json_data.get('company_email', ''),
            '{{conpany_web}}': json_data.get('conpany_web', ''),
            '{{table_2}}': json_data.get('table_2', ''),
            '{{checkbox_1}}': json_data.get('checkbox_1', ''),
            '{{checkbox_2}}': json_data.get('checkbox_2', ''),
            '{{created_no}}': json_data.get('created_no', ''),
            '{{created_day}}': json_data.get('created_day', ''),
            '{{created_month}}': json_data.get('created_month', ''),
            '{{created_year}}': json_data.get('created_year', ''),
            '{{checkbox_3}}': json_data.get('checkbox_3', ''),
            '{{checkbox_4}}': json_data.get('checkbox_4', ''),
            '{{own_name}}': json_data.get('own_name', ''),
            '{{own_sex}}': json_data.get('own_sex', ''),
            '{{dob_day}}': json_data.get('dob_day', ''),
            '{{dob_month}}': json_data.get('dob_month', ''),
            '{{dob_year}}': json_data.get('dob_year', ''),
            '{{kind}}': json_data.get('kind', ''),
            '{{national}}': json_data.get('national', ''),
            '{{table_4_1}}': json_data.get('table_4_1', ''),
            '{{id_no}}': json_data.get('id_no', ''),
            '{{id_created_day}}': json_data.get('id_created_day', ''),
            '{{id_created_month}}': json_data.get('id_created_month', ''),
            '{{id_created_year}}': json_data.get('id_created_year', ''),
            '{{id_created_place}}': json_data.get('id_created_place', ''),
            '{{id_expired_day}}': json_data.get('id_expired_day', ''),
            '{{id_expired_month}}': json_data.get('id_expired_month', ''),
            '{{id_expired_year}}': json_data.get('id_expired_year', ''),
            '{{live_address}}': json_data.get('live_address', ''),
            '{{live_ward}}': json_data.get('live_ward', ''),
            '{{live_province}}': json_data.get('live_province', ''),
            '{{live_city}}': json_data.get('live_city', ''),
            '{{live_nation}}': json_data.get('live_nation', ''),
            '{{contact_address}}': json_data.get('contact_address', ''),
            '{{contact_ward}}': json_data.get('contact_ward', ''),
            '{{contact_district}}': json_data.get('contact_district', ''),
            '{{contact_city}}': json_data.get('contact_city', ''),
            '{{contact_nation}}': json_data.get('contact_nation', ''),
            '{{contact_phone}}': json_data.get('contact_phone', ''),
            '{{contact_email}}': json_data.get('contact_email', ''),
            '{{pro_no}}': json_data.get('pro_no', ''),
            '{{pro_date}}': json_data.get('pro_date', ''),
            '{{pro_month}}': json_data.get('pro_month', ''),
            '{{pro_year}}': json_data.get('pro_year', ''),
            '{{pro_gov}}': json_data.get('pro_gov', ''),
            '{{coop_name}}': json_data.get('coop_name', ''),
            '{{coop_no}}': json_data.get('coop_no', ''),
            '{{coop_created_date}}': json_data.get('coop_created_date', ''),
            '{{coop_created_month}}': json_data.get('coop_created_month', ''),
            '{{coop_created_year}}': json_data.get('coop_created_year', ''),
            '{{coop_created_place}}': json_data.get('coop_created_place', ''),
            '{{coop_address}}': json_data.get('coop_address', ''),
            '{{coop_ward}}': json_data.get('coop_ward', ''),
            '{{coop_district}}': json_data.get('coop_district', ''),
            '{{coop_city}}': json_data.get('coop_city', ''),
            '{{coop_nation}}': json_data.get('coop_nation', ''),
            '{{coop_phone}}': json_data.get('coop_phone', ''),
            '{{coop_fax}}': json_data.get('coop_fax', ''),
            '{{coop_email}}': json_data.get('coop_email', ''),
            '{{coop_web}}': json_data.get('coop_web', ''),
            '{{pro_coop_no}}': json_data.get('pro_coop_no', ''),
            '{{pro_coop_date}}': json_data.get('pro_coop_date', ''),
            '{{pro_coop_month}}': json_data.get('pro_coop_month', ''),
            '{{pro_coop_year}}': json_data.get('pro_coop_year', ''),
            '{{pro_coop_gov}}': json_data.get('pro_coop_gov', ''),
            '{{table_6}}': json_data.get('table_6', ''),
            '{{amount}}': json_data.get('amount', ''),
            '{{amount_word}}': json_data.get('amount_word', ''),
            '{{amount_foreign}}': json_data.get('amount_foreign', ''),
            '{{checkbox_7}}': json_data.get('checkbox_7', ''),
            '{{checkbox_8}}': json_data.get('checkbox_8', ''),
            '{{gov_money}}': json_data.get('gov_money', ''),
            '{{gov_per}}': json_data.get('gov_per', ''),
            '{{per_money}}': json_data.get('per_money', ''),
            '{{per_per}}': json_data.get('per_per', ''),
            '{{for_money}}': json_data.get('for_money', ''),
            '{{for_per}}': json_data.get('for_per', ''),
            '{{oth_money}}': json_data.get('oth_money', ''),
            '{{oth_per}}': json_data.get('oth_per', ''),
            '{{sum}}': json_data.get('sum', ''),
            '{{sumPer}}': json_data.get('sumPer', ''),
            '{{vnd}}': json_data.get('vnd', ''),
            '{{vndPer}}': json_data.get('vndPer', ''),
            '{{usd}}': json_data.get('usd', ''),
            '{{usdPer}}': json_data.get('usdPer', ''),
            '{{gold}}': json_data.get('gold', ''),
            '{{goldPer}}': json_data.get('goldPer', ''),
            '{{land}}': json_data.get('land', ''),
            '{{landPer}}': json_data.get('landPer', ''),
            '{{mind}}': json_data.get('mind', ''),
            '{{mindPer}}': json_data.get('mindPer', ''),
            '{{other}}': json_data.get('other', ''),
            '{{otherPer}}': json_data.get('otherPer', ''),
            '{{proSum}}': json_data.get('proSum', ''),
            '{{proPer}}': json_data.get('proPer', ''),
            '{{preName}}': json_data.get('preName', ''),
            '{{preSex}}': json_data.get('preSex', ''),
            '{{preTitle}}': json_data.get('preTitle', ''),
            '{{preDobDate}}': json_data.get('preDobDate', ''),
            '{{preDobMonth}}': json_data.get('preDobMonth', ''),
            '{{preDobYear}}': json_data.get('preDobYear', ''),
            '{{preKind}}': json_data.get('preKind', ''),
            '{{preNation}}': json_data.get('preNation', ''),
            '{{table_4_2}}': json_data.get('table_4_2', ''),
            '{{preIdNo}}': json_data.get('preIdNo', ''),
            '{{preIdDate}}': json_data.get('preIdDate', ''),
            '{{preIdMonth}}': json_data.get('preIdMonth', ''),
            '{{preIdYear}}': json_data.get('preIdYear', ''),
            '{{preIdPlace}}': json_data.get('preIdPlace', ''),
            '{{preIdExDate}}': json_data.get('preIdExDate', ''),
            '{{preIdExMonth}}': json_data.get('preIdExMonth', ''),
            '{{preIdExYear}}': json_data.get('preIdExYear', ''),
            '{{preAdd}}': json_data.get('preAdd', ''),
            '{{preWard}}': json_data.get('preWard', ''),
            '{{preDist}}': json_data.get('preDist', ''),
            '{{preCity}}': json_data.get('preCity', ''),
            '{{preLiveNation}}': json_data.get('preLiveNation', ''),
            '{{preContactAdd}}': json_data.get('preContactAdd', ''),
            '{{preContactWard}}': json_data.get('preContactWard', ''),
            '{{preContactDist}}': json_data.get('preContactDist', ''),
            '{{preContactCity}}': json_data.get('preContactCity', ''),
            '{{preContactNation}}': json_data.get('preContactNation', ''),
            '{{prePhone}}': json_data.get('prePhone', ''),
            '{{preEmail}}': json_data.get('preEmail', ''),
            '{{ceo_name}}': json_data.get('ceo_name', ''),
            '{{ceo_phone}}': json_data.get('ceo_phone', ''),
            '{{account_name}}': json_data.get('account_name', ''),
            '{{account_phone}}': json_data.get('account_phone', ''),
            '{{tax_address}}': json_data.get('tax_address', ''),
            '{{tax_ward}}': json_data.get('tax_ward', ''),
            '{{tax_district}}': json_data.get('tax_district', ''),
            '{{tax_city}}': json_data.get('tax_city', ''),
            '{{tax_phone}}': json_data.get('tax_phone', ''),
            '{{tax_fax}}': json_data.get('tax_fax', ''),
            '{{tax_email}}': json_data.get('tax_email', ''),
            '{{start_date}}': json_data.get('start_date', ''),
            '{{start_month}}': json_data.get('start_month', ''),
            '{{start_year}}': json_data.get('start_year', ''),
            '{{checkbox_9}}': json_data.get('checkbox_9', ''),
            '{{checkbox_10}}': json_data.get('checkbox_10', ''),
            '{{checkbox_11}}': json_data.get('checkbox_11', ''),
            '{{fi_start_date}}': json_data.get('fi_start_date', ''),
            '{{fi_start_month}}': json_data.get('fi_start_month', ''),
            '{{fi_end_date}}': json_data.get('fi_end_date', ''),
            '{{fi_end_month}}': json_data.get('fi_end_month', ''),
            '{{employer_count}}': json_data.get('employer_count', ''),
            '{{checkbox_12}}': json_data.get('checkbox_12', ''),
            '{{checkbox_13}}': json_data.get('checkbox_13', ''),
            '{{table_5}}': json_data.get('table_5', ''),
            '{{checkbox_14}}': json_data.get('checkbox_14', ''),
            '{{checkbox_15}}': json_data.get('checkbox_15', ''),
            '{{checkbox_16}}': json_data.get('checkbox_16', ''),
            '{{checkbox_17}}': json_data.get('checkbox_17', ''),
            '{{table_7}}': json_data.get('table_7', ''),
            '{{transName}}': json_data.get('transName', ''),
            '{{transID}}': json_data.get('transID', ''),
            '{{transNo}}': json_data.get('transNo', ''),
            '{{transNoDate}}': json_data.get('transNoDate', ''),
            '{{transNoMonth}}': json_data.get('transNoMonth', ''),
            '{{transNoYear}}': json_data.get('transNoYear', ''),
            '{{transNoPlace}}': json_data.get('transNoPlace', ''),
            '{{famName}}': json_data.get('famName', ''),
            '{{famID}}': json_data.get('famID', ''),
            '{{famIdDate}}': json_data.get('famIdDate', ''),
            '{{famIdMonth}}': json_data.get('famIdMonth', ''),
            '{{famIdYear}}': json_data.get('famIdYear', ''),
            '{{famIdPlace}}': json_data.get('famIdPlace', ''),
            '{{famNo}}': json_data.get('famNo', ''),
            '{{famAdd}}': json_data.get('famAdd', ''),
            '{{famOwnName}}': json_data.get('famOwnName', ''),
            '{{table_4_3}}': json_data.get('table_4_3', ''),
            '{{famOwnId}}': json_data.get('famOwnId', ''),
            '{{famOwnIdDate}}': json_data.get('famOwnIdDate', ''),
            '{{famOwnIdMonth}}': json_data.get('famOwnIdMonth', ''),
            '{{famOwnIdYear}}': json_data.get('famOwnIdYear', ''),
            '{{famOwnIdPlace}}': json_data.get('famOwnIdPlace', ''),
            '{{famOwnIdExDate}}': json_data.get('famOwnIdExDate', ''),
            '{{famOwnIdExMonth}}': json_data.get('famOwnIdExMonth', ''),
            '{{famOwnIdExYear}}': json_data.get('famOwnIdExYear', ''),
            '{{SocialName}}': json_data.get('SocialName', ''),
            '{{SocialNo}}': json_data.get('SocialNo', ''),
            '{{SocialNoDate}}': json_data.get('SocialNoDate', ''),
            '{{SocialNoMonth}}': json_data.get('SocialNoMonth', ''),
            '{{SocialNoYear}}': json_data.get('SocialNoYear', ''),
            '{{SocialNoPlace}}': json_data.get('SocialNoPlace', ''),
            '{{SocialNoMST}}': json_data.get('SocialNoMST', ''),
            '{{SocialAdd}}': json_data.get('SocialAdd', ''),
            '{{SocialPreName}}': json_data.get('SocialPreName', ''),
            '{{table_4_4}}': json_data.get('table_4_4', ''),
            '{{SocialPreID}}': json_data.get('SocialPreID', ''),
            '{{SocialPreIdDate}}': json_data.get('SocialPreIdDate', ''),
            '{{SocialPreIdMonth}}': json_data.get('SocialPreIdMonth', ''),
            '{{SocialPreIdYear}}': json_data.get('SocialPreIdYear', ''),
            '{{SocialPreIdPlace}}': json_data.get('SocialPreIdPlace', ''),
            '{{SocialPreIdExDate}}': json_data.get('SocialPreIdExDate', ''),
            '{{SocialPreIdExMonth}}': json_data.get('SocialPreIdExMonth', ''),
            '{{SocialPreIdExYear}}': json_data.get('SocialPreIdExYear', ''),

        }
        
        table_data = json_data.get('tableData', [])
        
        # Gọi hàm fill_data với các giá trị cần thiết
        fill_data('TNHH1TV.docx', data, 'tnhh1tv_x.pdf',table_data)
        
        return JsonResponse({'status': 'success', 'message': 'Document created successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
import io
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
from django.views.decorators.http import require_POST
# TNHH1TV

chk_check = '☒'
chk_null = '☐'
chk_size = 28
font_size = 14
width_Chk = Pt(12)
width_txt = Pt(50)

dict_tables = [
              # 1   _0
              [
                'Thành lập mới',
                'Thành lập trên cơ sở tách doanh nghiệp',
                'Thành lập trên cơ sở chia doanh nghiệp',
                'Thành lập trên cơ sở hợp nhất doanh nghiệp',
                'Thành lập trên cơ sở chuyển đổi loại hình doanh nghiệp',
                'Thành lập trên cơ sở chuyển đổi từ hộ kinh doanh',
                'Thành lập trên cơ sở chuyển đổi từ cơ sở bảo trợ xã hội/quỹ xã hội/quỹ từ thiện'],
              # 3  _1
              [
                'Khu công nghiệp',
                'Khu chế xuất',
                'Khu kinh tế',
                'Khu công nghệ cao'],
              # 5, 9, 14, 15    _2
              [
                'Chứng minh nhân dân',
                'Căn cước công dân',
                'Hộ chiếu',
                'Loại khác (ghi rõ):……'],
              # 5 Mô hình tổ chức công ty:  _3
              [
                'Hội đồng thành viên, Giám đốc hoặc Tổng Giám đốc',
                'Chủ tịch công ty, Giám đốc hoặc Tổng Giám đốc'],

              # 10.9  _4
              [
                'Khấu trừ',
                'Trực tiếp trên GTGT',
                'Trực tiếp trên doanh số',
                'Không phải nộp thuế GTGT'],
              # 11  _5
              [
                'Tự in hóa đơn',
                'Đặt in hóa đơn',
                'Sử dụng hóa đơn điện tử',
                'Mua hóa đơn của cơ quan thuế'],
              # 12  _6
              [
                'Hàng tháng',
                '03 tháng một lần',
                '06 tháng một lần'],
              # 10_5  _7
              [
                'Hạch toán độc lập',
                'Hạch toán phụ thuộc',
                'Có báo cáo tài chính hợp nhất'],
              # 10_8  _8
              [
                'Có',
                'Không']
              ]
# print(dict_tables["table_1"])

def set_border_cell(cell):
    # https://stackoverflow.com/questions/76790690/python-docx-adding-shapes-rectangle
    cell = cell
    tc_pr = cell._element.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    tc_pr.append(borders)

    for border_type in ['top', 'left', 'bottom', 'right']:
        border_elm = OxmlElement(f'w:{border_type}')
        border_elm.set(qn('w:val'), 'single')
        border_elm.set(qn('w:sz'), '6')
        border_elm.set(qn('w:space'), '0')
        border_elm.set(qn('w:color'), '000000')
        borders.append(border_elm)

def set_cell_text_with_font(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
      if(paragraph.text == chk_check or paragraph.text == chk_null):
        for run in paragraph.runs:
            run.font.size = Pt(chk_size)
      else:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)

def fill_data(doc_path, data, output_path,table_data):
    
    # Open the DOCX file
    doc = Document(doc_path)
    print('1')
    # Replace placeholders with data
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
                paragraph.style.font.size = Pt(font_size)
    print('2')
    # print(table_data)

    # 1
    set_checkbox_list(doc, 0, dict_tables[0], data.get('{{table_1}}'))
    # 1 ------------------------------------------------------------------------

    # 3
    set_checkbox_list(doc, 1, dict_tables[1], data.get('{{table_2}}'))
    # 3 ------------------------------------------------------------------------

    # 4
    # delete rows null
    # doc.tables[2].rows[1]._element.getparent().remove(doc.tables[2].rows[1]._element)

    # add row with data
    print('3')
    if(table_data != ""):
      # chạy sau header => i+1
      row =1
      for table in table_data:
        j=0
        # rows (0,0), (0,1), (0,2),...
        for cellValues in table:
          txt = str(cellValues)
          # check vs uncheck
          if(txt == 'checked'):
            set_cell_text_with_font(doc.tables[2].cell(row,j), chk_check)
          elif(txt == 'unchecked'):
            set_cell_text_with_font(doc.tables[2].cell(row,j), chk_null)
          else:
            set_cell_text_with_font(doc.tables[2].cell(row,j), txt)
          set_border_cell(doc.tables[2].cell(row,j))
          doc.tables[2].cell(row,j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
          j = j+1
        # add row
        doc.tables[2].add_row()
        row = row+1
    # 4 ------------------------------------------------------------------------
    print('4')
    # 5
    set_checkbox_table2x2(doc, 3, 1, [data.get('{{table_4_1}}')])

    # b) Đối với chủ sở hữu là tổ chức:
    set_checkbox_list(doc, 4, dict_tables[3], data.get('{{table_6}}'))
    # 5 ------------------------------------------------------------------------
    print('5')

    # 7
    
    set_cell_text_with_font(doc.tables[5].cell(1,1) ,data.get('{{gov_money}}') )
    set_cell_text_with_font(doc.tables[5].cell(1,2) ,data.get('{{gov_per}}') )
    set_cell_text_with_font(doc.tables[5].cell(2,1) ,data.get('{{per_money}}') )
    set_cell_text_with_font(doc.tables[5].cell(2,2) ,data.get('{{per_per}}') )
    set_cell_text_with_font(doc.tables[5].cell(3,1) ,data.get('{{for_money}}') )
    set_cell_text_with_font(doc.tables[5].cell(3,2) ,data.get('{{for_per}}') )
    set_cell_text_with_font(doc.tables[5].cell(4,1) ,data.get('{{oth_money}}') )
    set_cell_text_with_font(doc.tables[5].cell(4,2) ,data.get('{{oth_per}}') )
    set_cell_text_with_font(doc.tables[5].cell(5,1) ,data.get('{{sum}}') )
    set_cell_text_with_font(doc.tables[5].cell(5,2) ,data.get('{{sumPer}}') )
    
    # doc.tables[5].cell(1,1).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # doc.tables[5].cell(1,2).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # doc.tables[5].cell(2,1).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # doc.tables[5].cell(2,2).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # doc.tables[5].cell(3,1).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # doc.tables[5].cell(3,2).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # doc.tables[5].cell(4,1).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # doc.tables[5].cell(4,2).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # doc.tables[5].cell(5,1).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # doc.tables[5].cell(5,2).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # 7 ------------------------------------------------------------------------

     # 8

    set_cell_text_with_font(doc.tables[6].cell(1,2) , data.get('{{vnd}}'))
    # doc.tables[6].cell(1,2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(1,3) , data.get('{{vndPer}}'))
    # doc.tables[6].cell(1,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(2,2) , data.get('{{usd}}'))
    # doc.tables[6].cell(2,2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(2,3) , data.get('{{usdPer}}'))
    # doc.tables[6].cell(2,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(3,2) , data.get('{{gold}}'))
    # doc.tables[6].cell(3,2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(3,3) , data.get('{{goldPer}}'))
    # doc.tables[6].cell(3,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(4,2) , data.get('{{land}}'))
    # doc.tables[6].cell(4,2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(4,3) , data.get('{{landPer}}'))
    # doc.tables[6].cell(4,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(5,2) , data.get('{{mind}}'))
    # doc.tables[6].cell(5,2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(5,3) , data.get('{{mindPer}}'))
    # doc.tables[6].cell(5,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(6,2) , data.get('{{other}}'))
    # doc.tables[6].cell(6,2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(6,3) , data.get('{{otherPer}}'))
    # doc.tables[6].cell(6,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(7,2) , data.get('{{proSum}}'))
    # doc.tables[6].cell(7,2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[6].cell(7,3) , data.get('{{proPer}}'))
    # doc.tables[6].cell(7,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 8 ------------------------------------------------------------------------
    # 9
    set_checkbox_table2x2(doc, 7, 1, [data.get('{{table_4_2}}')])
    # doc.tables[7]
    # 9 ------------------------------------------------------------------------
    
# 10
    str10_1 = "Thông tin về Giám đốc/Tổng giám đốc (nếu có): \n Họ và tên Giám đốc/Tổng giám đốc:  "+ data.get('{{ceo_name}}')+"\n Điện thoại: "+ data.get('{{ceo_phone}}')
    str10_2 = "Thông tin về Kế toán trưởng/Phụ trách kế toán (nếu có): \n Họ và tên Kế toán trưởng/Phụ trách kế toán: "+ data.get('{{account_name}}')+" \n Điện thoại: "+ data.get('{{account_phone}}')
    str10_3 = "Địa chỉ nhận thông báo thuế (chỉ kê khai nếu địa chỉ nhận thông báo thuế khác địa chỉ trụ sở chính): \n Số nhà, ngách, hẻm, ngõ, đường phố/tổ/xóm/ấp/thôn:"+ data.get('{{tax_address}}')+" \n Xã/Phường/Thị trấn: "+ data.get('{{tax_ward}}')+" \n Quận/Huyện/Thị xã/Thành phố thuộc tỉnh: "+ data.get('{{tax_district}}')+" \n Tỉnh/Thành phố: "+ data.get('{{tax_city}}')+"\n Điện thoại (nếu có): "+ data.get('{{tax_phone}}')+" Fax (nếu có): "+ data.get('{{tax_fax}}')+"\nEmail (nếu có): "+ data.get('{{tax_email}}')
    str10_4 = "Ngày bắt đầu hoạt động7 (trường hợp doanh nghiệp dự kiến bắt đầu hoạt động kể từ ngày được cấp Giấy chứng nhận đăng ký doanh nghiệp thì không cần kê khai nội dung này): \n "+ data.get('{{start_date}}')+"/"+ data.get('{{start_month}}')+"/"+ data.get('{{start_year}}')
    str10_5 = []
    str10_6 = "Năm tài chính: Áp dụng từ ngày "+ data.get('{{fi_start_date}}')+"/"+ data.get('{{fi_start_month}}')+" đến ngày "+ data.get('{{fi_end_date}}')+"/"+ data.get('{{fi_end_month}}')+" \n (ghi ngày, tháng bắt đầu và kết thúc niên độ kế toán)"
    str10_7 = "Tổng số lao động (dự kiến): "+ data.get('{{employer_count}}')
    str10_8 =  data.get('{{checkbox_12}}')
    str10_8_2 =  data.get('{{checkbox_13}}')
    str10_9 =  data.get('{{table_5}}')
    print(str10_1)
    # 10.1
    set_cell_text_with_font(doc.tables[8].cell(1,1) , str10_1)
    # doc.tables[8].cell(1,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # 10.2
    set_cell_text_with_font(doc.tables[8].cell(2,1) , str10_2)
    # doc.tables[8].cell(2,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # 10.3
    set_cell_text_with_font(doc.tables[8].cell(3,1) , str10_3)
    # doc.tables[8].cell(3,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # 10.4
    set_cell_text_with_font(doc.tables[8].cell(4,1) , str10_4)
    # doc.tables[8].cell(4,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # 10.5
    set_cell_text_with_font(doc.tables[8].cell(5,1) , "Hình thức hạch toán (Đánh dấu X vào một trong hai ô “Hạch toán độc lập” hoặc “Hạch toán phụ thuộc”. Trường hợp tích chọn ô “Hạch toán độc lập” mà thuộc đối tượng phải lập và gửi báo cáo tài chính hợp nhất cho cơ quan có thẩm quyền theo quy định thì tích chọn thêm ô “Có báo cáo tài chính hợp nhất”): ")
    # doc.tables[8].cell(5,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.tables[8].cell(5,1).add_table(1,1)
    doc.tables[8].cell(5,1).tables[0].columns[0].width = 3*width_txt
    set_cell_text_with_font(doc.tables[8].cell(5,1).tables[0].cell(0,0) , "Hạch toán độc lập")
    # doc.tables[8].cell(5,1).tables[0].cell(0,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.tables[8].cell(5,1).tables[0].add_column((width_txt))
    set_cell_text_with_font(doc.tables[8].cell(5,1).tables[0].cell(0,1) , chk_null)
    doc.tables[8].cell(5,1).tables[0].add_column(3*(width_txt))
    set_cell_text_with_font(doc.tables[8].cell(5,1).tables[0].cell(0,2) , "Có báo cáo tài chính hợp nhất")
    # doc.tables[8].cell(5,1).tables[0].cell(0,2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.tables[8].cell(5,1).tables[0].add_column((width_txt))
    set_cell_text_with_font(doc.tables[8].cell(5,1).tables[0].cell(0,3) , chk_null)
    doc.tables[8].cell(5,1).tables[0].add_row()
    set_cell_text_with_font(doc.tables[8].cell(5,1).tables[0].cell(1,0) , "Hạch toán phụ thuộc")
    # doc.tables[8].cell(5,1).tables[0].cell(1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_text_with_font(doc.tables[8].cell(5,1).tables[0].cell(1,1) , chk_null)

    if(data.get('{{checkbox_9}}')== dict_tables[7][0]):
      set_cell_text_with_font(doc.tables[8].cell(5,1).tables[0].cell(0,1) , chk_check)
    if(data.get('{{checkbox_10}}')== dict_tables[7][1]):
      set_cell_text_with_font(doc.tables[8].cell(5,1).tables[0].cell(1,1) , chk_check)
    if(data.get('{{checkbox_11}}')== dict_tables[7][2]):
      set_cell_text_with_font(doc.tables[8].cell(5,1).tables[0].cell(0,3) , chk_check)
    # 10.6
    set_cell_text_with_font(doc.tables[8].cell(6,1) , str10_6)
    # doc.tables[8].cell(6,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 10.7
    set_cell_text_with_font(doc.tables[8].cell(7,1) , str10_7)
    # doc.tables[8].cell(7,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 10.8
    set_cell_text_with_font(doc.tables[8].cell(8,1) , "Hoạt động theo dự án BOT/BTO/BT/BOO, BLT, BTL, O&M:")
    doc.tables[8].cell(8,1).add_table(1,1)
    doc.tables[8].cell(8,1).tables[0].columns[0].width = 2*width_txt
    set_cell_text_with_font(doc.tables[8].cell(8,1).tables[0].cell(0,0) , dict_tables[8][0])
    # doc.tables[8].cell(8,1).tables[0].cell(0,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.tables[8].cell(8,1).tables[0].add_column((width_txt))
    set_cell_text_with_font(doc.tables[8].cell(8,1).tables[0].cell(0,1) , chk_null)
    doc.tables[8].cell(8,1).tables[0].add_column(2*(width_txt))
    set_cell_text_with_font(doc.tables[8].cell(8,1).tables[0].cell(0,2) , dict_tables[8][1])
    # doc.tables[8].cell(8,1).tables[0].cell(0,2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.tables[8].cell(8,1).tables[0].add_column((width_txt))
    set_cell_text_with_font(doc.tables[8].cell(8,1).tables[0].cell(0,3) , chk_null)

    # check select
    # for chk in dict_tables[8]:
    #   if(chk == str10_8):
    #     position = dict_tables[8].index(chk)
    #     if(position == 0):
    #       set_cell_text_with_font(doc.tables[8].cell(8,1).tables[0].cell(0,1) , chk_check)
    #     if(position == 1):
    #       set_cell_text_with_font(doc.tables[8].cell(8,1).tables[0].cell(0,3) , chk_check)
    if(str10_8 == dict_tables[8][0]):
      set_cell_text_with_font(doc.tables[8].cell(8,1).tables[0].cell(0,1) , chk_check)
    if(str10_8_2 == dict_tables[8][1]):
      set_cell_text_with_font(doc.tables[8].cell(8,1).tables[0].cell(0,3) , chk_check)
    

    # 10.9
    set_cell_text_with_font(doc.tables[8].cell(9,5) , chk_null)
    set_cell_text_with_font(doc.tables[8].cell(9,8) , chk_null)
    set_cell_text_with_font(doc.tables[8].cell(9,11) , chk_null)
    set_cell_text_with_font(doc.tables[8].cell(9,14) , chk_null)

    # check select
    for chk in dict_tables[4]:
      if(chk == str10_9):
        position = dict_tables[4].index(chk)
        if(position == 0):
          set_cell_text_with_font(doc.tables[8].cell(9,5), chk_check)
        if(position == 1):
          set_cell_text_with_font(doc.tables[8].cell(9,8), chk_check)
        if(position == 2):
          set_cell_text_with_font(doc.tables[8].cell(9,11), chk_check)
        if(position == 3):
          set_cell_text_with_font(doc.tables[8].cell(9,14), chk_check)

    # print(doc.tables[8].cell(9,4).text)
    # khau tu
    # 10 -----------------------------------------------------------------------
    print('6')
    # 11
    str_bills = [data.get('{{checkbox_14}}'), data.get('{{checkbox_15}}'), data.get('{{checkbox_16}}'), data.get('{{checkbox_17}}')]
    set_checkbox_table2x2(doc, 9, 2, str_bills)
# 11 -----------------------------------------------------------------------

    # 12
    # (0,0) checkbox (0,1) Hàng tháng (0,2) checkbox (0,3) 03 tháng một lần (0,4) checkbox (0,5) 06 tháng một lần
    #
    set_cell_text_with_font(doc.tables[10].cell(0,0), chk_null)
    doc.tables[10].cell(0,0).width = width_Chk
    set_cell_text_with_font(doc.tables[10].cell(0,1), dict_tables[6][0])
    doc.tables[10].cell(0,1).width = width_txt
    # doc.tables[10].cell(0,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    #
    set_cell_text_with_font(doc.tables[10].cell(0,2), chk_null)
    doc.tables[10].cell(0,2).width = width_Chk
    doc.tables[10].add_column(width_txt)
    set_cell_text_with_font(doc.tables[10].cell(0,3), dict_tables[6][1])
    # doc.tables[10].cell(0,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    #
    doc.tables[10].add_column(width_Chk)
    set_cell_text_with_font(doc.tables[10].cell(0,4), chk_null)
    doc.tables[10].cell(0,4).width = width_Chk
    doc.tables[10].add_column(width_txt)
    set_cell_text_with_font(doc.tables[10].cell(0,5), dict_tables[6][2])
    # doc.tables[10].cell(0,5).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # check select
    for chk in dict_tables[6]:
      if(chk == data.get('{{table_7}}')):
        position = dict_tables[6].index(chk)
        if(position == 0):
          set_cell_text_with_font(doc.tables[10].cell(0,0), chk_check)
        if(position == 1):
          set_cell_text_with_font(doc.tables[10].cell(0,2), chk_check)
        if(position == 2):
          set_cell_text_with_font(doc.tables[10].cell(0,4), chk_check)

    # print(doc.tables[10].cell(0,0).text)
    # 12 -----------------------------------------------------------------------
    print('7')
    # 14
    set_checkbox_table2x2(doc, 11, 1, [data.get('{{table_4_3}}')])
    # print(doc.tables[11].cell(0,0).text)
    # 14 -----------------------------------------------------------------------
    print('8')
    # 15
    set_checkbox_table2x2(doc, 12, 1, [data.get('{{table_4_4}}')])
    # print(doc.tables[12].cell(0,0).text)
    # 15 -----------------------------------------------------------------------
    
    # set_cell_text_with_font(doc.tables[0].cell(0,1),'⊠', 40)
    # if doc.tables:
    #   table = doc.tables[1]
    #   for row_idx, row in enumerate(table.rows):
    #     for col_idx, cell in enumerate(row.cells):
    #         print(f"Cell at row {row_idx}, column {col_idx}: {cell.text}")

    # if doc.tables:
    #   tables = doc.tables
    #   for table in tables:
    #     # print(table)
    #     for row in table.rows:
    #       # print(row)
    #       for cell in row.cells:
    #         # print(cell)
    #         for paragraph in cell.paragraphs:
    #           # if(paragraph.text == "Thành lập mới"):
    #             print(paragraph.text)


    # Save the modified document
    print('chuẩn bị xuất')
    doc.save(output_path)
    print('đã xuất')

def set_checkbox_list(doc, index_table, array, text):
    # index_table: thứ tự bảng trong file
    # array: mảng mẫu
    # text: giá trị chọn
    val = text
    i=0
    for chk in array:
      if(chk == val):
        position = array.index(chk)
        set_cell_text_with_font(doc.tables[index_table].cell(position,1), chk_check)
      else:
        set_cell_text_with_font(doc.tables[index_table].cell(i,1), chk_null)
      i = i+1

def set_checkbox_table2x2(doc, index_table, style, check_index):
    # index_table: thứ tự bảng trong file
    # style: loại bảng 1 or 2
    # check_index: vị trí check

    # style 1: 5 9 14 15
    # (0,0) checkbox (0,1) cmdd (0,2) checkbox (0,3) cccd
    # (1,0) checkbox (1,1) hc (1,2) checkbox (1,3) !=

    # style 2: 11
    # (0,0) checkbox (0,1) Tự in hóa đơn (0,2) checkbox (0,3) Đặt in hóa đơn
    # (1,0) checkbox (1,1) Sử dụng hóa đơn điện tử (1,2) checkbox (1,3) Mua hóa đơn của cơ quan thuế

    str_id = dict_tables[2]
    str_bills = dict_tables[5]
    str = []
    if (style == 1):
      str = str_id
    elif (style == 2):
      str = str_bills
    # checkbox
    # thông số
    position_table = index_table

    #
    set_cell_text_with_font(doc.tables[position_table].cell(0,0), chk_null)
    set_cell_text_with_font(doc.tables[position_table].cell(1,0), chk_null)
    # set width
    doc.tables[position_table].cell(0,0).width = width_Chk
    doc.tables[position_table].cell(1,0).width = width_Chk

    set_cell_text_with_font(doc.tables[position_table].cell(0,1), str[0])
    set_cell_text_with_font(doc.tables[position_table].cell(1,1), str[2])
    # set width
    doc.tables[position_table].cell(0,1).width = width_txt
    doc.tables[position_table].cell(1,1).width = width_txt
    # set vertical_alignment
    doc.tables[position_table].cell(0,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.tables[position_table].cell(1,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # add column
    doc.tables[position_table].add_column(width_Chk)
    set_cell_text_with_font(doc.tables[position_table].cell(0,2), chk_null)
    set_cell_text_with_font(doc.tables[position_table].cell(1,2), chk_null)
    doc.tables[position_table].add_column(width_txt)
    set_cell_text_with_font(doc.tables[position_table].cell(0,3), str[1])
    set_cell_text_with_font(doc.tables[position_table].cell(1,3), str[3])
    # set vertical_alignment
    doc.tables[position_table].cell(0,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.tables[position_table].cell(1,3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # check box select
    for chk in str:
      for arr in check_index:
        if(arr == chk):
          position = str.index(arr)
          if(position == 0):
            set_cell_text_with_font(doc.tables[position_table].cell(0,0), chk_check)
          if(position == 1):
            set_cell_text_with_font(doc.tables[position_table].cell(0,2), chk_check)
          if(position == 2):
            set_cell_text_with_font(doc.tables[position_table].cell(1,0), chk_check)
          if(position == 3):
            set_cell_text_with_font(doc.tables[position_table].cell(1,2), chk_check)



# Data to fill in the DOCX
# data = {
    
# }

# table_data = [
#     ["1", "test", "t", ""],
#     ["Row2, Col1", "Row2, Col2", "Row2, Col3", "Row2, Col4"],
#     ["Row3, Col1", "Row3, Col2", "Row3, Col3", "Row3, Col4"],
#     ["Row4, Col1", "Row4, Col2", "Row4, Col3", "Row4, Col4"],
#     ["Row5, Col1", "Row5, Col2", "Row5, Col3", "Row5, Col4"],
#     ["Row6, Col1", "Row6, Col2", "Row6, Col3", "Row6, Col4"]
# ]
# table_data = json_data.get('table_data', [])


# Fill the data into the document
# fill_data('TNHH1TV.docx', data, 'tnhh1tv_xuat.docx',table_data)

